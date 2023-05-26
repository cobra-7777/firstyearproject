from docx import Document
from datetime import date

doc_name = "Lecture " + date.today()

def text_to_doc(title, input_text): 
    retries = 0
    success = False
    while retries <= 5 and success == False:
        try:
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(input_text)
            doc.save(doc_name)
            success = True
        except:
            if retries >= 5:
                print("Critical error, unable to create a doc file. Stopping program...")
            else:
                retries = retries + 1
                print("Failed to create a doc file. Retrying for the " + str(retries) + " time...")